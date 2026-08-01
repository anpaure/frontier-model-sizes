from __future__ import annotations

import csv
import math
import os
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


DOCUMENTS_CACHE_ROOT = (
    Path.home() / ".codex/plugins/cache/openai-primary-runtime/documents"
)
DOCUMENT_SCRIPT_DIRS = sorted(
    (
        candidate
        for candidate in DOCUMENTS_CACHE_ROOT.glob("*/skills/documents/scripts")
        if (candidate / "table_geometry.py").is_file()
    ),
    reverse=True,
)
if not DOCUMENT_SCRIPT_DIRS:
    raise RuntimeError(
        "Could not locate the bundled documents table_geometry.py helper under "
        f"{DOCUMENTS_CACHE_ROOT}"
    )
sys.path.insert(0, str(DOCUMENT_SCRIPT_DIRS[0]))
from table_geometry import apply_table_geometry  # noqa: E402
from normalize_ooxml_zip import normalize_ooxml_zip  # noqa: E402


WORK_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = WORK_DIR / "outputs/019f6c42-2d53-7743-ab07-6293e2618dd7"
FORECAST_LEDGER_PATH = WORK_DIR / "sources/human_parameter_forecasts_2026-07-17.csv"
REGISTRY_VERSION = "2.1"
REGISTRY_DATE = "2026-07-17"
OUTPUT_PATH = OUTPUT_DIR / f"frontier_parameter_prediction_registry_v{REGISTRY_VERSION}_2026-07-17.docx"
PRIMARY_MODELS = ("Claude Fable 5", "GPT-5.6 Sol")


# Resolved compact_reference_guide preset.
FONT = "Calibri"
INK = "243838"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667274"
GRID = "D2DCDC"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def add_page_field(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def set_cell_text(cell, text, *, bold=False, size=9.5, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    set_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(
    table,
    widths,
    *,
    numeric_cols=(),
    header_size=9.0,
    body_size=9.5,
    cell_margins_dxa=None,
):
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=cell_margins_dxa or CELL_MARGINS,
    )
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_cant_split(row)
    for c, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(
            cell,
            cell.text,
            bold=True,
            size=header_size,
            color=DARK_BLUE,
            align=WD_ALIGN_PARAGRAPH.CENTER if c in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT,
        )
    for row in table.rows[1:]:
        for c, cell in enumerate(row.cells):
            set_cell_text(
                cell,
                cell.text,
                size=body_size,
                align=WD_ALIGN_PARAGRAPH.CENTER if c in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT,
            )


def add_para(doc, text="", *, size=11, bold=False, color=INK, after=6, before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_label_value(doc, label, value, *, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.5, color=INK, bold=True)
    r = p.add_run(value)
    set_font(r, size=10.5, color=INK)
    return p


def add_lead_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(f"{label}  ")
    set_font(r, size=11, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_font(r, size=11, color=INK)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        s = styles[style_name]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    title_style = styles.add_style("Registry Title", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = FONT
    title_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(INK)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(4)

    subtitle_style = styles.add_style("Registry Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = FONT
    subtitle_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle_style.font.size = Pt(11)
    subtitle_style.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(14)

    source_style = styles.add_style("Table Source", WD_STYLE_TYPE.PARAGRAPH)
    source_style.font.name = FONT
    source_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    source_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    source_style.font.size = Pt(9)
    source_style.font.italic = True
    source_style.font.color.rgb = RGBColor.from_string(MUTED)
    source_style.paragraph_format.space_before = Pt(4)
    source_style.paragraph_format.space_after = Pt(4)


def add_header_footer(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("Frontier Parameter Forecast Registry")
    set_font(r, size=9, color=MUTED, bold=True)
    r = hp.add_run("    |    Living record")
    set_font(r, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)


def geometric_midpoint(low, high):
    return math.sqrt(low * high)


def geometric_mean(values):
    if not values:
        raise ValueError("Cannot aggregate an empty forecast set")
    return math.exp(sum(math.log(value) for value in values) / len(values))


def load_forecasts():
    with FORECAST_LEDGER_PATH.open(newline="", encoding="utf-8-sig") as handle:
        rows = list(csv.DictReader(handle))
    required = {
        "forecast_id",
        "contributor",
        "date",
        "model",
        "forecast_text",
        "low_t",
        "high_t",
        "central_t",
        "confidence",
        "provenance",
        "notes",
        "supersedes",
    }
    if not rows or set(rows[0]) != required:
        raise ValueError(f"Forecast ledger schema mismatch: {FORECAST_LEDGER_PATH}")
    ids = [row["forecast_id"] for row in rows]
    if len(ids) != len(set(ids)):
        raise ValueError("Forecast ledger contains duplicate forecast_id values")
    id_set = set(ids)
    superseded_ids = set()
    for row in rows:
        low = float(row["low_t"])
        high = float(row["high_t"])
        if not row["forecast_id"] or not row["contributor"] or not row["model"]:
            raise ValueError(f"Missing required identity field in {row}")
        if low <= 0 or high <= 0 or low > high:
            raise ValueError(f"Invalid bounds in {row['forecast_id']}: {low}, {high}")
        row["low_t"] = low
        row["high_t"] = high
        central_text = row["central_t"].strip()
        central = float(central_text) if central_text else None
        if central is not None and (central < low or central > high):
            raise ValueError(f"Stated central outside bounds in {row['forecast_id']}: {central}")
        row["central_t"] = central
        row["point_t"] = central if central is not None else geometric_midpoint(low, high)
        if row["supersedes"]:
            if row["supersedes"] not in id_set:
                raise ValueError(f"Unknown supersedes id in {row['forecast_id']}")
            superseded_ids.add(row["supersedes"])
    active = [row for row in rows if row["forecast_id"] not in superseded_ids]
    keys = [(row["contributor"], row["model"]) for row in active]
    if len(keys) != len(set(keys)):
        raise ValueError("Multiple active forecasts exist for the same contributor/model pair")
    return rows, active


def crowd_stats(active, model):
    rows = [row for row in active if row["model"] == model]
    points = [row["point_t"] for row in rows]
    overlap_low = max(row["low_t"] for row in rows)
    overlap_high = min(row["high_t"] for row in rows)
    return {
        "rows": rows,
        "n": len(rows),
        "center": geometric_mean(points),
        "envelope_low": min(row["low_t"] for row in rows),
        "envelope_high": max(row["high_t"] for row in rows),
        "overlap_low": overlap_low,
        "overlap_high": overlap_high,
        "point_span": max(points) / min(points),
    }


def displayed_forecast(row):
    text = row["forecast_text"]
    confidence = row["confidence"].strip()
    if confidence and confidence.lower() != "not specified" and confidence not in text:
        text = f"{text} ({confidence})"
    return text


def short_model(model):
    return {"Claude Fable 5": "Fable", "GPT-5.6 Sol": "Sol"}.get(model, model)


def main():
    _, active = load_forecasts()
    stats = {model: crowd_stats(active, model) for model in PRIMARY_MODELS}
    respondent_order = list(dict.fromkeys(row["contributor"] for row in active))
    primary_respondents = [
        respondent
        for respondent in respondent_order
        if any(row["contributor"] == respondent and row["model"] in PRIMARY_MODELS for row in active)
    ]
    respondent_count = len(set(row["contributor"] for row in active))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    add_header_footer(section)
    configure_styles(doc)

    kicker = add_para(doc, "LIVING REGISTRY", size=9.5, bold=True, color=BLUE, after=3)
    kicker.paragraph_format.keep_with_next = True
    doc.add_paragraph("Frontier Model Parameter Prediction Registry", style="Registry Title")
    doc.add_paragraph(
        "Human forecasts with provenance and crowd aggregation, separate from model-generated estimates.",
        style="Registry Subtitle",
    )
    add_label_value(doc, "Version", f"{REGISTRY_VERSION} · 17 July 2026")
    add_label_value(doc, "Scope", f"Total base parameters in T · Anonymous respondents: {respondent_count}")

    add_lead_callout(
        doc,
        "Current crowd signal",
        f"Equal-weight geometric aggregation gives {stats['Claude Fable 5']['center']:.1f}T for Claude Fable 5 "
        f"(n={stats['Claude Fable 5']['n']}) and {stats['GPT-5.6 Sol']['center']:.1f}T for GPT-5.6 Sol "
        f"(n={stats['GPT-5.6 Sol']['n']}). These are tracking statistics rather than calibrated posterior estimates.",
    )

    doc.add_heading("1. Anonymous forecast register", level=1)
    add_para(
        doc,
        "Forecasts are preserved as stated under stable anonymous respondent IDs. An explicit central estimate is used when supplied; otherwise range-only forecasts use geometric midpoints for multiplicative uncertainty. Revisions append rather than overwrite, and no name-to-ID mapping is stored in the project.",
        after=7,
    )

    table = doc.add_table(rows=1, cols=7)
    headers = [
        "Respondent ID",
        "Date",
        "Fable forecast",
        "Sol forecast",
        "Fable point",
        "Sol point",
        "Provenance",
    ]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    rows = []
    for respondent in primary_respondents:
        respondent_rows = [row for row in active if row["contributor"] == respondent]
        by_model = {row["model"]: row for row in respondent_rows}
        fable = by_model.get("Claude Fable 5")
        sol = by_model.get("GPT-5.6 Sol")
        dates = list(dict.fromkeys(row["date"] for row in respondent_rows if row["model"] in PRIMARY_MODELS))
        provenance = list(dict.fromkeys(row["provenance"] for row in respondent_rows))
        confidences = list(
            dict.fromkeys(
                row["confidence"]
                for row in respondent_rows
                if row["confidence"] and row["confidence"].lower() != "not specified"
            )
        )
        provenance_text = " ".join(provenance)
        if confidences:
            provenance_text += f" Confidence: {', '.join(confidences)}."
        rows.append(
            [
                respondent,
                " / ".join(dates),
                displayed_forecast(fable) if fable else "—",
                displayed_forecast(sol) if sol else "—",
                f"{fable['point_t']:.1f}T" if fable else "—",
                f"{sol['point_t']:.1f}T" if sol else "—",
                provenance_text,
            ]
        )
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    style_table(
        table,
        [1500, 900, 950, 950, 850, 850, 3360],
        numeric_cols=(1, 2, 3, 4, 5),
        header_size=8.6,
        body_size=9.2,
        cell_margins_dxa={"top": 40, "bottom": 40, "start": 120, "end": 120},
    )
    source = doc.add_paragraph(style="Table Source")
    source.add_run(f"Source: {FORECAST_LEDGER_PATH.name}; user-supplied statements in the Codex conversation dated 17 July 2026.")

    other_label = add_para(
        doc,
        "Additional model forecasts",
        size=11,
        bold=True,
        color=DARK_BLUE,
        before=5,
        after=4,
    )
    other_label.paragraph_format.keep_with_next = True
    other_rows = [row for row in active if row["model"] not in PRIMARY_MODELS]
    other_by_respondent = {}
    for row in other_rows:
        other_by_respondent.setdefault(row["contributor"], []).append(row)
    for respondent, records in other_by_respondent.items():
        forecasts = "; ".join(f"{row['model']} — {displayed_forecast(row)}" for row in records)
        provenance = " ".join(dict.fromkeys(row["provenance"] for row in records))
        add_para(
            doc,
            f"{respondent} ({records[0]['date']}; {provenance}): {forecasts}.",
            size=9.5,
            after=4,
        )
    doc.add_heading("2. Current crowd aggregation", level=1)
    crowd = doc.add_table(rows=1, cols=6)
    for i, text in enumerate(
        ["Model", "n", "Crowd center", "Range envelope", "Range overlap", "Disagreement diagnostic"]
    ):
        crowd.rows[0].cells[i].text = text
    crowd_rows = []
    for model in PRIMARY_MODELS:
        model_stats = stats[model]
        overlap = (
            f"{model_stats['overlap_low']:.1f}-{model_stats['overlap_high']:.1f}T"
            if model_stats["overlap_low"] <= model_stats["overlap_high"]
            else "None"
        )
        crowd_rows.append(
            [
                model,
                str(model_stats["n"]),
                f"{model_stats['center']:.1f}T",
                f"{model_stats['envelope_low']:.1f}-{model_stats['envelope_high']:.1f}T",
                overlap,
                f"Individual centers span {model_stats['point_span']:.1f}x; the crowd mean hides substantial disagreement.",
            ]
        )
    for values in crowd_rows:
        cells = crowd.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    style_table(
        crowd,
        [1650, 500, 1250, 1350, 1250, 3360],
        numeric_cols=(1, 2, 3, 4),
        header_size=8.8,
        body_size=9.3,
    )

    doc.add_heading("3. Aggregation policy", level=1)
    policy_items = [
        (
            "Raw record",
            "Store anonymous respondent ID, date/time, model/version, forecast, provenance class, confidence, evidence basis, and revision link.",
        ),
        (
            "Point conversion",
            "Use a respondent's explicit central estimate when supplied; otherwise convert [L, H] to sqrt(L x H), reflecting multiplicative uncertainty.",
        ),
        (
            "Crowd center",
            "Use the equal-weight geometric mean unless weights were set before outcomes became known.",
        ),
        (
            "Independence",
            "Flag shared evidence or discussion; correlated opinions are not independent evidence.",
        ),
        (
            "Revisions",
            "Append a new row and mark the prior entry superseded; never erase it.",
        ),
        (
            "Human versus AI",
            "Keep human and AI forecasts separate; label any combined ensemble as secondary.",
        ),
    ]
    policy_table = doc.add_table(rows=1, cols=2)
    policy_table.rows[0].cells[0].text = "Rule"
    policy_table.rows[0].cells[1].text = "Implementation"
    for label, detail in policy_items:
        cells = policy_table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    style_table(policy_table, [2100, 7260], header_size=9.0, body_size=9.2)

    doc.add_page_break()
    doc.add_heading("4. Revision log", level=1)
    revision = doc.add_table(rows=1, cols=3)
    revision.rows[0].cells[0].text = "Version"
    revision.rows[0].cells[1].text = "Date"
    revision.rows[0].cells[2].text = "Change"
    revision_rows = [
        ["1.0", "2026-07-17", "Created the registry and fixed the aggregation and provenance rules."],
        ["2.0", "2026-07-17", "Expanded the ledger to 14 respondents; retained equal-respondent geometric pooling."],
        [
            REGISTRY_VERSION,
            "2026-08-01",
            "Expanded the ledger to 21 respondents and replaced all respondent names, name-bearing record IDs, and personal provenance references with stable anonymous IDs; no name-to-ID mapping is retained.",
        ],
    ]
    current_summary = ", ".join(
        f"{short_model(model)} {stats[model]['center']:.1f}T (n={stats[model]['n']})" for model in PRIMARY_MODELS
    )
    revision_rows[-1][2] += f" Current crowds: {current_summary}."
    for values in revision_rows:
        cells = revision.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    style_table(
        revision,
        [1100, 1500, 6760],
        numeric_cols=(0, 1),
        header_size=9.0,
        body_size=9.0,
        cell_margins_dxa={"top": 40, "bottom": 40, "start": 120, "end": 120},
    )

    props = doc.core_properties
    props.title = "Frontier Model Parameter Prediction Registry"
    props.subject = "Living record of anonymous human parameter-count forecasts for Claude Fable 5 and GPT-5.6 Sol"
    props.author = "Project team"
    props.last_modified_by = "Project team"
    props.keywords = "frontier models, parameter estimates, prediction registry, wisdom of crowds"
    props.comments = f"Version {REGISTRY_VERSION} updated 17 July 2026; generated from {FORECAST_LEDGER_PATH.name}"

    doc.save(OUTPUT_PATH)
    normalize_ooxml_zip(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
