from __future__ import annotations

import csv
import json
import math
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs/019f6c42-2d53-7743-ab07-6293e2618dd7"
LEDGER = ROOT / "sources/human_parameter_forecasts_2026-07-17.csv"
REGISTRY = OUT / "frontier_parameter_prediction_registry_v2.1_2026-07-17.docx"
WORKBOOK = OUT / "frontier_parameter_model_crowd_50pct_2026-07-17.xlsx"
TARGETS = ("Claude Fable 5", "GPT-5.6 Sol")
OPENROUTER_RAW = ROOT / "sources/openrouter_operational_snapshot_2026-07-18.json.gz"
OPENROUTER_MODELS = ROOT / "sources/openrouter_model_signals_2026-07-18.csv"
OPENROUTER_PROVIDERS = ROOT / "sources/openrouter_provider_signals_2026-07-18.csv"
OPENROUTER_TIERS = ROOT / "sources/openrouter_endpoint_tier_signals_2026-07-18.csv"
OPENROUTER_DAILY = ROOT / "sources/openrouter_throughput_daily_2026-07-18.csv"
OPENROUTER_MODEL_HISTORY = ROOT / "sources/openrouter_model_snapshot_history_2026-07-18.csv"
OPENROUTER_PROVIDER_HISTORY = ROOT / "sources/openrouter_provider_snapshot_history_2026-07-18.csv"
OPENROUTER_TIER_HISTORY = ROOT / "sources/openrouter_endpoint_tier_snapshot_history_2026-07-18.csv"
OPENROUTER_DAILY_HISTORY = ROOT / "sources/openrouter_throughput_daily_history_2026-07-18.csv"
OPENROUTER_HISTORY_MANIFEST = ROOT / "sources/openrouter_snapshot_history_manifest_2026-07-18.csv"
OPENROUTER_AUDIT = OUT / "openrouter_epoch_match_audit_2026-07-18.csv"
OPENROUTER_RESULT = OUT / "openrouter_parameter_signal_backtest_2026-07-18.json"
OPENROUTER_TEMPORAL_RESULT = OUT / "openrouter_temporal_stability_audit_2026-07-18.json"
OPENROUTER_ENDPOINT_STABILITY = OUT / "openrouter_endpoint_temporal_stability_2026-07-18.csv"
OPENROUTER_MODEL_STABILITY = OUT / "openrouter_model_temporal_stability_2026-07-18.csv"
OPENROUTER_REFRESH_STABILITY = OUT / "openrouter_refresh_stability_2026-07-18.csv"
OPENROUTER_TIER_PREDICTIONS = OUT / "openrouter_tier_counterfactual_predictions_2026-07-18.csv"
OPENROUTER_COLLECTION_AUDIT = OUT / "openrouter_collection_audit_2026-07-18.json"
OPENROUTER_OFFICIAL_SNAPSHOT = ROOT / "sources/openrouter_official_endpoint_snapshot_2026-07-18.json.gz"
OPENROUTER_OFFICIAL_PRICES = ROOT / "sources/openrouter_official_endpoint_prices_2026-07-18.csv"
OPENROUTER_OFFICIAL_COMPARISON = OUT / "openrouter_official_endpoint_crosscheck_2026-07-18.csv"
OPENROUTER_OFFICIAL_AUDIT = OUT / "openrouter_official_endpoint_audit_2026-07-18.json"
ECI_REPRODUCED_SCORES = ROOT / "sources/epoch_eci_reproduced_scores_2026-07-31.csv"
ECI_REPRODUCTION_METADATA = ROOT / "sources/epoch_eci_reproduction_metadata_2026-07-31.json"
ECI_REPRODUCTION_CROSSCHECK = OUT / "epoch_eci_reproduction_crosscheck_2026-07-31.csv"
ECI_REPRODUCTION_AUDIT = OUT / "epoch_eci_reproduction_audit_2026-07-31.json"
AA_EXPANDED_RESULT = OUT / "aa_expanded_parameter_audit_2026-07-18.json"
AA_EXPANDED_PANEL = OUT / "aa_expanded_parameter_panel_2026-07-18.csv"
AA_EXPANDED_PREDICTIONS = OUT / "aa_expanded_parameter_predictions_2026-07-18.csv"
AA_EXPANDED_OVERLAPS = OUT / "aa_expanded_parameter_overlap_audit_2026-07-18.csv"
AA_DETAILED_RAW = ROOT / "sources/aa_detailed_snapshot_2026-07-31.html.gz"
AA_DETAILED_MODELS = ROOT / "sources/aa_detailed_model_signals_2026-07-31.csv"
AA_DETAILED_METADATA = ROOT / "sources/aa_detailed_collection_metadata_2026-07-31.json"
AA_CALIBRATION_OVERRIDES = ROOT / "sources/aa_calibration_primary_overrides_2026-07-31.json"
AA_PARAMETER_LABEL_AVAILABILITY = ROOT / "sources/aa_parameter_label_availability_2026-07-31.json"
AA_SCORE_AVAILABILITY = ROOT / "sources/aa_score_availability_2026-07-31.json"
AA_CHANGELOG_RAW = ROOT / "sources/aa_changelog_2026-07-31.json.gz"
AA_SCORE_TIMING_AUDIT = OUT / "aa_score_availability_timing_audit_2026-07-31.json"
AA_SCORE_TIMING_CHANGES = OUT / "aa_score_availability_timing_changes_2026-07-31.csv"
OPEN_MODEL_PARAMETER_TRUTH = ROOT / "sources/open_model_parameter_truth_reconciliation_2026-07-31.json"
ECI_ARCHITECTURE_BLEND_RESULT = OUT / "eci_architecture_blend_challenger_2026-07-31.json"
ECI_ARCHITECTURE_BLEND_PREDICTIONS = OUT / "eci_architecture_blend_challenger_predictions_2026-07-31.csv"
AA_INFERENCE_RESULT = OUT / "aa_inference_budget_audit_2026-07-18.json"
AA_DETAILED_PANEL = OUT / "aa_detailed_parameter_panel_2026-07-18.csv"
AA_REASONING_PAIRS = OUT / "aa_reasoning_pair_audit_2026-07-18.csv"
AA_DETAILED_CROSSCHECK = OUT / "aa_detailed_epoch_crosscheck_2026-07-18.csv"
AA_INFERENCE_PREDICTIONS = OUT / "aa_inference_budget_predictions_2026-07-18.csv"
AA_OPERATIONAL_RESULT = OUT / "aa_operational_signal_audit_2026-07-18.json"
AA_OPERATIONAL_PANEL = OUT / "aa_operational_parameter_panel_2026-07-18.csv"
AA_OPERATIONAL_PREDICTIONS = OUT / "aa_operational_backtest_predictions_2026-07-18.csv"
AA_OPENROUTER_CROSSCHECK = OUT / "aa_openrouter_operational_crosscheck_2026-07-18.csv"
ECI_COMPONENT_RESULT = OUT / "eci_component_extended_audit_2026-07-18.json"
ECI_COMPONENT_PANEL = OUT / "eci_component_expanded_parameter_panel_2026-07-18.csv"
ECI_COMPONENT_COMPARISON = OUT / "eci_component_active_incremental_comparison_2026-07-18.csv"
ECI_MULTIVARIATE_RESULT = OUT / "eci_multivariate_component_audit_2026-07-18.json"
ECI_MULTIVARIATE_PREDICTIONS = OUT / "eci_multivariate_component_predictions_2026-07-18.csv"
ECI_MULTIVARIATE_NARROW_CI_PREDICTIONS = OUT / "eci_multivariate_component_narrow_eci_ci_predictions_2026-07-18.csv"
ECI_MULTIVARIATE_TARGETS = OUT / "eci_multivariate_component_targets_2026-07-18.csv"
ECI_MULTIVARIATE_COVERAGE = OUT / "eci_multivariate_component_coverage_2026-07-18.csv"
POSTTRAINING_LINEAGE_RESULT = OUT / "posttraining_lineage_audit_2026-07-18.json"
POSTTRAINING_LINEAGE_EDGES = OUT / "posttraining_lineage_edges_2026-07-18.csv"
POSTTRAINING_LINEAGE_MEASUREMENTS = OUT / "posttraining_lineage_measurements_2026-07-18.csv"
POSTTRAINING_LINEAGE_PREDICTIONS = OUT / "posttraining_lineage_predictions_2026-07-18.csv"
FRONTIER_SHARED_BASE_SENSITIVITY = OUT / "frontier_shared_base_sensitivity_2026-07-18.csv"
FRONTIER_LINEAGE_EVIDENCE = OUT / "frontier_lineage_evidence_2026-07-18.csv"
ACTIVE_TRANSPORT_RESULT = OUT / "active_parameter_transport_audit_2026-07-18.json"
ACTIVE_TRANSPORT_PREDICTIONS = OUT / "active_parameter_transport_predictions_2026-07-18.csv"
ACTIVE_TRANSPORT_TARGETS = OUT / "active_parameter_transport_targets_2026-07-18.csv"
K3_ARCHITECTURE_FACTS = ROOT / "sources/kimi_k3_release_evidence_2026-07-31.json"
OPENROUTER_ACTIVE_PRICE_RESULT = OUT / "openrouter_active_price_audit_2026-07-18.json"
OPENROUTER_ACTIVE_PRICE_MATCHES = OUT / "openrouter_active_parameter_match_audit_2026-07-18.csv"
OPENROUTER_ACTIVE_PRICE_PREDICTIONS = OUT / "openrouter_active_price_predictions_2026-07-18.csv"
OPENROUTER_ACTIVE_PRICE_TARGETS = OUT / "openrouter_active_price_targets_2026-07-18.csv"
NO_COT_EXACT_DATE_AUDIT = OUT / "no_cot_exact_date_audit_2026-07-18.json"
NO_COT_EXACT_DATE_MODELS = OUT / "no_cot_exact_date_model_audit_2026-07-18.csv"
NO_COT_ARCHITECTURE_AUDIT = OUT / "no_cot_architecture_elasticity_audit_2026-07-18.json"
NO_COT_ARCHITECTURE_PREDICTIONS = OUT / "no_cot_architecture_elasticity_predictions_2026-07-18.csv"
NO_COT_EXACT_DATE_OVERRIDES = ROOT / "sources/no_cot_exact_date_overrides_2026-07-18.csv"
NO_COT_EXACT_DATE_METADATA = ROOT / "sources/no_cot_exact_date_collection_metadata_2026-07-18.json"
QWEN_EXACT_DATE_RAW = ROOT / "sources/qwen3_30b_a3b_instruct_2507_hf_commits_2026-07-18.json.gz"
FRONTIER_PRIMARY_EVIDENCE = ROOT / "sources/frontier_primary_evidence_2026-07-18.csv"
FRONTIER_PRIMARY_METADATA = ROOT / "sources/frontier_primary_evidence_collection_metadata_2026-07-18.json"
FRONTIER_PRIMARY_CLAIMS = ROOT / "sources/anthropic_fable_mythos_primary_claims_2026-07-18.json"
FRONTIER_PRIMARY_OPENAI_RAW = ROOT / "sources/openai_gpt_5_6_system_card_2026-07-18.html.gz"
FRONTIER_PRIMARY_AUDIT = OUT / "frontier_primary_evidence_audit_2026-07-18.json"
FRONTIER_PRIMARY_CONTROLS = OUT / "frontier_primary_evidence_controls_2026-07-18.csv"
METR_PRIMARY_SIGNALS = ROOT / "sources/metr_horizon_official_signals_2026-07-18.csv"
METR_PRIMARY_RAW = ROOT / "sources/metr_benchmark_results_1_1_2026-07-18.yaml"
METR_PRIMARY_METADATA = ROOT / "sources/metr_horizon_official_metadata_2026-07-18.json"
METR_PRIMARY_AUDIT = OUT / "metr_primary_source_audit_2026-07-18.json"
METR_LEGACY_CROSSCHECK = ROOT / "sources/metr_horizon_user_snapshot_2026-07-17.csv"
IKP_AUDIT = OUT / "ikp_parameter_signal_audit_2026-07-18.json"
IKP_PREDICTIONS = OUT / "ikp_parameter_chronological_predictions_2026-07-18.csv"
IKP_OVERLAP = OUT / "ikp_parameter_incremental_overlap_2026-07-18.csv"
IKP_CONDITIONAL_AUDIT = OUT / "ikp_conditional_benchmark_signal_audit_2026-07-18.json"
IKP_CONDITIONAL_PREDICTIONS = OUT / "ikp_conditional_benchmark_predictions_2026-07-18.csv"
IKP_SOURCE_METADATA = ROOT / "sources/ikp_source_metadata_2026-07-18.json"
INDEPENDENT_REAUDIT = OUT / "codex_independent_reaudit_metrics_2026-07-17.json"
EPOCH_FEEDBACK_SOURCE = ROOT / "sources/epoch_employee_calibration_feedback_2026-07-21.csv"
EPOCH_FEEDBACK_AUDIT = OUT / "epoch_feedback_lean_architecture_audit_2026-07-21.json"
EPOCH_FEEDBACK_PANEL = OUT / "epoch_feedback_critique_panel_2026-07-21.csv"
EPOCH_FEEDBACK_PREDICTIONS = OUT / "lean_architecture_predictions_2026-07-21.csv"
EPOCH_FEEDBACK_TARGETS = OUT / "lean_architecture_target_sensitivity_2026-07-21.csv"
OPUS5_EVIDENCE = ROOT / "sources/claude_opus_5_evidence_2026-07-31.json"


def geomean(values):
    return math.exp(sum(math.log(value) for value in values) / len(values))


with LEDGER.open(newline="", encoding="utf-8-sig") as handle:
    rows = list(csv.DictReader(handle))

ids = [row["forecast_id"] for row in rows]
assert len(ids) == len(set(ids)), "duplicate forecast_id"
id_set = set(ids)
superseded_ids = {row["supersedes"] for row in rows if row["supersedes"]}
assert superseded_ids <= id_set, "unknown supersedes id"
active = [row for row in rows if row["forecast_id"] not in superseded_ids]
active_keys = [(row["contributor"], row["model"]) for row in active]
assert len(active_keys) == len(set(active_keys)), "multiple active contributor/model forecasts"


def crowd_point(row):
    low = float(row["low_t"])
    high = float(row["high_t"])
    central_text = (row.get("central_t") or "").strip()
    central = float(central_text) if central_text else None
    assert central is None or low <= central <= high, f"central outside bounds: {row['forecast_id']}"
    return central if central is not None else math.sqrt(low * high)


r16 = {(row["model"], crowd_point(row)): row for row in active if row["contributor"] == "Respondent R16"}
assert ("Claude Fable 5", 6.4) in r16 and ("GPT-5.6 Sol", 5.2) in r16
r21 = {(row["model"], row["low_t"], row["high_t"]): row for row in active if row["contributor"] == "Respondent R21"}
assert ("Claude Fable 5", "5.0", "10.0") in r21
assert ("GPT-5.6 Sol", "3.0", "8.0") in r21
r17_opus = next(row for row in active if row["contributor"] == "Respondent R17" and row["model"] == "Claude Opus 4.7 / 4.8 shared base")
assert math.isclose(crowd_point(r17_opus), math.sqrt(0.8 * 2.0), rel_tol=0, abs_tol=1e-12)

centers = {}
counts = {}
for model in TARGETS:
    model_rows = [row for row in active if row["model"] == model]
    assert model_rows, f"missing required crowd pool: {model}"
    counts[model] = len(model_rows)
    points = [crowd_point(row) for row in model_rows]
    centers[model] = geomean(points)

doc = Document(REGISTRY)
registry_rows = [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]
contributors = {row["contributor"] for row in active}
assert len(registry_rows) == len(contributors) + 1
assert [row[0] for row in registry_rows[1:]] == list(dict.fromkeys(row["contributor"] for row in active))
r16_row = next(row for row in registry_rows[1:] if row[0] == "Respondent R16")
assert r16_row[2] == "6.4T (range 5.0-7.0T)" and r16_row[3] == "5.2T (range 4.0-6.0T)"
assert r16_row[4] == "6.4T" and r16_row[5] == "5.2T"
r21_row = next(row for row in registry_rows[1:] if row[0] == "Respondent R21")
assert r21_row[2] == "5.0-10.0T" and r21_row[3] == "3.0-8.0T"
assert r21_row[4] == "7.1T" and r21_row[5] == "4.9T"
document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
assert f"Anonymous respondents: {len(contributors)}" in document_text

crowd_rows = [[cell.text.strip() for cell in row.cells] for row in doc.tables[1].rows]
crowd_by_model = {row[0]: row for row in crowd_rows[1:]}
for model in TARGETS:
    assert crowd_by_model[model][1] == str(counts[model])
    assert crowd_by_model[model][2] == f"{centers[model]:.1f}T"

workbook = load_workbook(WORKBOOK, data_only=True, read_only=True)
posterior = workbook["Horizon Estimates"]
final = workbook["Final Ensemble"]
posterior_headers = {cell.value: index + 1 for index, cell in enumerate(posterior[5])}
final_headers = {cell.value: index + 1 for index, cell in enumerate(final[5])}

for row_number in range(6, 16):
    model = final.cell(row_number, final_headers["Model / base"]).value
    if model not in TARGETS:
        continue
    workbook_crowd = float(final.cell(row_number, final_headers["Crowd center (T)"]).value)
    evidence_model = float(final.cell(row_number, final_headers["Evidence model (T)"]).value)
    final_forecast = float(final.cell(row_number, final_headers["Final forecast (T)"]).value)
    assert math.isclose(workbook_crowd, centers[model], rel_tol=0, abs_tol=1e-10)
    assert math.isclose(final_forecast, math.sqrt(evidence_model * centers[model]), rel_tol=0, abs_tol=1e-9)

for row_number in range(6, 16):
    model = posterior.cell(row_number, posterior_headers["Model / base"]).value
    if model in TARGETS:
        value = float(posterior.cell(row_number, posterior_headers["Crowd point (T)"]).value)
        assert math.isclose(value, centers[model], rel_tol=0, abs_tol=1e-10)

fable_posterior_row = next(
    row_number
    for row_number in range(6, 16)
    if posterior.cell(row_number, posterior_headers["Model / base"]).value == "Claude Fable 5"
)
baseline_fable = (
    float(posterior.cell(fable_posterior_row, posterior_headers["Existing central (T)"]).value)
    ** float(posterior.cell(fable_posterior_row, posterior_headers["Existing weight"]).value)
    * float(posterior.cell(fable_posterior_row, posterior_headers["Horizon prior (T)"]).value)
    ** float(posterior.cell(fable_posterior_row, posterior_headers["Horizon weight"]).value)
    * float(posterior.cell(fable_posterior_row, posterior_headers["Compute-structured prior (T)"]).value)
    ** float(posterior.cell(fable_posterior_row, posterior_headers["Compute weight"]).value)
)
ikp_result_for_posterior = json.loads(IKP_AUDIT.read_text(encoding="utf-8"))
ikp_fable_t = (
    ikp_result_for_posterior["target_signal"]["fable"]["strict_open_only_release_and_vendor_holdout"]
    ["mean"]["estimates"]["forward_inverse"]["estimated_b"] / 1000
)
ikp_evidence_weight = ikp_result_for_posterior["decision"]["incremental_evidence_weight"]
expected_fable_evidence = (
    baseline_fable ** (1 - ikp_evidence_weight)
    * ikp_fable_t**ikp_evidence_weight
)
observed_fable_evidence = float(
    posterior.cell(fable_posterior_row, posterior_headers["Posterior (T)"]).value
)
assert math.isclose(observed_fable_evidence, expected_fable_evidence, rel_tol=0, abs_tol=1e-9)

opus_row = next(row_number for row_number in range(6, 16) if posterior.cell(row_number, posterior_headers["Model / base"]).value == "Claude Opus 4.7 / 4.8 shared base")
opus_crowd = float(posterior.cell(opus_row, posterior_headers["Crowd point (T)"]).value)
assert math.isclose(opus_crowd, math.sqrt(0.8 * 2.0), rel_tol=0, abs_tol=1e-12)
opus_final_row = next(row_number for row_number in range(6, 16) if final.cell(row_number, final_headers["Model / base"]).value == "Claude Opus 4.7 / 4.8 shared base")
assert final.cell(opus_final_row, final_headers["Crowd center (T)"]).value is None
assert final.cell(opus_final_row, final_headers["Crowd weight"]).value == 0

opus5_evidence = json.loads(OPUS5_EVIDENCE.read_text(encoding="utf-8"))
opus5_posterior_row = next(
    row_number for row_number in range(6, 16)
    if posterior.cell(row_number, posterior_headers["Model / base"]).value == "Claude Opus 5"
)
opus5_final_row = next(
    row_number for row_number in range(6, 16)
    if final.cell(row_number, final_headers["Model / base"]).value == "Claude Opus 5"
)
assert posterior.cell(opus5_posterior_row, posterior_headers["Horizon evidence"]).value == "Neutral — no direct horizon"
assert posterior.cell(opus5_posterior_row, posterior_headers["Crowd point (T)"]).value is None
assert final.cell(opus5_final_row, final_headers["Crowd center (T)"]).value is None
assert final.cell(opus5_final_row, final_headers["Crowd weight"]).value == 0
assert final.cell(opus5_final_row, final_headers["Final status"]).value == "Model posterior only"
assert 2.9 < float(final.cell(opus5_final_row, final_headers["Final forecast (T)"]).value) < 3.2
assert opus5_evidence["identity"]["base_identity_policy"] == "unique_base"
assert opus5_evidence["identity"]["parameter_disclosed"] is False
assert opus5_evidence["identity"]["same_weight_identity_disclosed"] is False
assert all(opus5_evidence["availability"][key] is False for key in ("metr", "no_cot", "ikp"))

manifest = workbook["Source Manifest"]
manifest_headers = [cell.value for cell in manifest[5]]
assert manifest_headers == ["Source", "Path", "Bytes", "SHA-256", "Snapshot date"]
manifest_snapshot_dates = {
    row[4]
    for row in manifest.iter_rows(min_row=6, values_only=True)
    if row[0]
}
assert manifest_snapshot_dates == {datetime(2026, 7, 31)}
def resolve_manifest_path(value):
    path = Path(str(value)).expanduser()
    if not path.is_absolute():
        path = ROOT / path
    return str(path.resolve())


manifest_paths = {
    resolve_manifest_path(row[1])
    for row in manifest.iter_rows(min_row=1, values_only=True)
    if len(row) > 1 and row[1]
}
assert str(LEDGER) in manifest_paths
assert str(REGISTRY) in manifest_paths
assert str(OPUS5_EVIDENCE) in manifest_paths
for relative_path in opus5_evidence["source_files"].values():
    assert str(ROOT / relative_path) in manifest_paths
for path in (OPENROUTER_RAW, OPENROUTER_MODELS, OPENROUTER_PROVIDERS, OPENROUTER_TIERS, OPENROUTER_DAILY, OPENROUTER_MODEL_HISTORY, OPENROUTER_PROVIDER_HISTORY, OPENROUTER_TIER_HISTORY, OPENROUTER_DAILY_HISTORY, OPENROUTER_HISTORY_MANIFEST, OPENROUTER_AUDIT, OPENROUTER_RESULT, OPENROUTER_TEMPORAL_RESULT, OPENROUTER_ENDPOINT_STABILITY, OPENROUTER_MODEL_STABILITY, OPENROUTER_REFRESH_STABILITY, OPENROUTER_TIER_PREDICTIONS, OPENROUTER_COLLECTION_AUDIT, OPENROUTER_OFFICIAL_SNAPSHOT, OPENROUTER_OFFICIAL_PRICES, OPENROUTER_OFFICIAL_COMPARISON, OPENROUTER_OFFICIAL_AUDIT):
    assert str(path) in manifest_paths
for path in (AA_EXPANDED_RESULT, AA_EXPANDED_PANEL, AA_EXPANDED_PREDICTIONS, AA_EXPANDED_OVERLAPS):
    assert str(path) in manifest_paths
for path in (AA_DETAILED_RAW, AA_DETAILED_MODELS, AA_DETAILED_METADATA, AA_CALIBRATION_OVERRIDES, AA_INFERENCE_RESULT, AA_DETAILED_PANEL, AA_REASONING_PAIRS, AA_DETAILED_CROSSCHECK, AA_INFERENCE_PREDICTIONS):
    assert str(path) in manifest_paths
aa_override_payload = json.loads(AA_CALIBRATION_OVERRIDES.read_text(encoding="utf-8"))
for override in aa_override_payload["overrides"]:
    for evidence in override["primary_source"]["local_evidence"]:
        assert str(ROOT / evidence["path"]) in manifest_paths
assert str(AA_PARAMETER_LABEL_AVAILABILITY) in manifest_paths
aa_label_timing_payload = json.loads(
    AA_PARAMETER_LABEL_AVAILABILITY.read_text(encoding="utf-8")
)
for record in aa_label_timing_payload["records"]:
    for evidence in record["local_evidence"]:
        assert str(ROOT / evidence["path"]) in manifest_paths
for path in (
    AA_SCORE_AVAILABILITY,
    AA_CHANGELOG_RAW,
    AA_SCORE_TIMING_AUDIT,
    AA_SCORE_TIMING_CHANGES,
    OPEN_MODEL_PARAMETER_TRUTH,
):
    assert str(path) in manifest_paths
parameter_truth_payload = json.loads(OPEN_MODEL_PARAMETER_TRUTH.read_text(encoding="utf-8"))
for source in parameter_truth_payload["source_files"]:
    assert str(ROOT / source["path"]) in manifest_paths
for path in (ECI_ARCHITECTURE_BLEND_RESULT, ECI_ARCHITECTURE_BLEND_PREDICTIONS):
    assert str(path) in manifest_paths
for path in (AA_OPERATIONAL_RESULT, AA_OPERATIONAL_PANEL, AA_OPERATIONAL_PREDICTIONS, AA_OPENROUTER_CROSSCHECK):
    assert str(path) in manifest_paths
for path in (ECI_COMPONENT_RESULT, ECI_COMPONENT_PANEL, ECI_COMPONENT_COMPARISON):
    assert str(path) in manifest_paths
for path in (ECI_REPRODUCED_SCORES, ECI_REPRODUCTION_METADATA, ECI_REPRODUCTION_CROSSCHECK, ECI_REPRODUCTION_AUDIT):
    assert str(path) in manifest_paths
for path in (ECI_MULTIVARIATE_RESULT, ECI_MULTIVARIATE_PREDICTIONS, ECI_MULTIVARIATE_NARROW_CI_PREDICTIONS, ECI_MULTIVARIATE_TARGETS, ECI_MULTIVARIATE_COVERAGE):
    assert str(path) in manifest_paths
for path in (POSTTRAINING_LINEAGE_RESULT, POSTTRAINING_LINEAGE_EDGES, POSTTRAINING_LINEAGE_MEASUREMENTS, POSTTRAINING_LINEAGE_PREDICTIONS, FRONTIER_SHARED_BASE_SENSITIVITY, FRONTIER_LINEAGE_EVIDENCE):
    assert str(path) in manifest_paths
for path in (ACTIVE_TRANSPORT_RESULT, ACTIVE_TRANSPORT_PREDICTIONS, ACTIVE_TRANSPORT_TARGETS, K3_ARCHITECTURE_FACTS):
    assert str(path) in manifest_paths
for path in (OPENROUTER_ACTIVE_PRICE_RESULT, OPENROUTER_ACTIVE_PRICE_MATCHES, OPENROUTER_ACTIVE_PRICE_PREDICTIONS, OPENROUTER_ACTIVE_PRICE_TARGETS):
    assert str(path) in manifest_paths
for path in (NO_COT_EXACT_DATE_AUDIT, NO_COT_EXACT_DATE_MODELS, NO_COT_EXACT_DATE_OVERRIDES, NO_COT_EXACT_DATE_METADATA, QWEN_EXACT_DATE_RAW):
    assert str(path) in manifest_paths
for path in (NO_COT_ARCHITECTURE_AUDIT, NO_COT_ARCHITECTURE_PREDICTIONS):
    assert str(path) in manifest_paths
for path in (FRONTIER_PRIMARY_EVIDENCE, FRONTIER_PRIMARY_METADATA, FRONTIER_PRIMARY_CLAIMS, FRONTIER_PRIMARY_OPENAI_RAW, FRONTIER_PRIMARY_AUDIT, FRONTIER_PRIMARY_CONTROLS):
    assert str(path) in manifest_paths
for path in (METR_PRIMARY_SIGNALS, METR_PRIMARY_RAW, METR_PRIMARY_METADATA, METR_PRIMARY_AUDIT, METR_LEGACY_CROSSCHECK):
    assert str(path) in manifest_paths
for path in (IKP_AUDIT, IKP_PREDICTIONS, IKP_OVERLAP, IKP_SOURCE_METADATA):
    assert str(path) in manifest_paths
for path in (EPOCH_FEEDBACK_SOURCE, EPOCH_FEEDBACK_AUDIT, EPOCH_FEEDBACK_PANEL, EPOCH_FEEDBACK_PREDICTIONS, EPOCH_FEEDBACK_TARGETS):
    assert str(path) in manifest_paths

epoch_feedback = json.loads(EPOCH_FEEDBACK_AUDIT.read_text(encoding="utf-8"))
assert EPOCH_FEEDBACK_SOURCE.exists()
assert EPOCH_FEEDBACK_PANEL.exists()
assert EPOCH_FEEDBACK_PREDICTIONS.exists()
assert EPOCH_FEEDBACK_TARGETS.exists()
assert epoch_feedback["feedback_reproduction"]["rows"] == 8
assert epoch_feedback["feedback_reproduction"]["all_displayed_formula_outputs_reproduced"] is True
assert epoch_feedback["feedback_reproduction"]["metrics"]["independent_base_clusters"] == 6
assert epoch_feedback["architecture_panel"]["rows"] == 89
assert epoch_feedback["architecture_panel"]["families"] == 40
assert epoch_feedback["promotion_gates"]["architecture_candidate"]["all_gates_pass"] is False
assert epoch_feedback["promotion_gates"]["active_sparsity_candidate"]["all_gates_pass"] is False
assert epoch_feedback["decision"]["change_live_model"] is False
assert epoch_feedback["decision"]["incremental_live_weight"] == 0

metr_result = json.loads(METR_PRIMARY_AUDIT.read_text(encoding="utf-8"))
assert metr_result["status"] == "PASS"
assert metr_result["official_asset"]["result_rows"] == 26
assert metr_result["losslessness"]["full_scaffold_entries"] == 114
assert metr_result["legacy_exact_crosscheck"]["exact_rows"] == 26
assert metr_result["legacy_exact_crosscheck"]["mismatch_count"] == 0

ikp_result = json.loads(IKP_AUDIT.read_text(encoding="utf-8"))
with IKP_OVERLAP.open(newline="", encoding="utf-8") as handle:
    ikp_overlap_rows = list(csv.DictReader(handle))
assert ikp_result["decision"]["promote_incremental_ikp_weight"] is False
assert ikp_result["decision"]["incremental_evidence_weight"] == 0
assert ikp_result["decision"]["incremental_final_weight_when_crowd_is_50pct"] == 0
assert ikp_result["incremental_overlap"]["models"] == len(ikp_overlap_rows)
assert ikp_result["incremental_overlap"]["families"] == len({row["family"] for row in ikp_overlap_rows})
assert (
    ikp_result["incremental_overlap"]["chronological_fixed_weight_subset"]["models"]
    < ikp_result["decision"]["evidence_gates"]["minimum_chronological_subset_models"]
)
assert ikp_result["incremental_overlap"]["chronological_fixed_weight_subset"]["family_bootstrap"]["ci_90"][1] < 0
ikp_conditional = json.loads(IKP_CONDITIONAL_AUDIT.read_text(encoding="utf-8"))
assert ikp_conditional["decision"]["conditional_incremental_signal_corroborated"] is True
assert ikp_conditional["decision"]["change_live_ikp_weight"] is False
assert ikp_conditional["heldout_results"]["gpqa_diamond"]["passing_specifications"] == 4
assert ikp_conditional["heldout_results"]["mmlu"]["passing_specifications"] == 3
assert ikp_conditional["upstream_reproduction"]["narrative_summary_audit"]["stale_claim_count"] == 6
independent_reaudit = json.loads(INDEPENDENT_REAUDIT.read_text(encoding="utf-8"))
ikp_reaudit = independent_reaudit["ikp_parameter_signal_reaudit"]
assert ikp_reaudit["all_pass"] is True
assert ikp_reaudit["published_fit_recomputation"]["configurations"] == 93
assert ikp_reaudit["serving_base_collapse"]["distinct_weight_bases"] == 87
assert ikp_reaudit["strict_fable_recomputation"]["anthropic_excluded"] is True
assert ikp_reaudit["incremental_overlap_recomputation"]["models"] == ikp_result["incremental_overlap"]["models"]
assert ikp_reaudit["incremental_overlap_recomputation"]["chronological_models"] == ikp_result["incremental_overlap"]["chronological_fixed_weight_subset"]["models"]
assert ikp_reaudit["workbook_integration"]["formula_absolute_difference"] < 1e-10
assert ikp_reaudit["workbook_integration"]["policy_absolute_difference"] < 1e-10
conditional_reaudit = independent_reaudit["ikp_conditional_benchmark_reaudit"]
assert conditional_reaudit["all_pass"] is True
assert conditional_reaudit["source_inventory"]["raw_configurations"] == 100
assert conditional_reaudit["source_inventory"]["benchmark_rows"] == 81
assert conditional_reaudit["source_inventory"]["weight_bases"] == 87
assert conditional_reaudit["prediction_integrity"]["rows"] == 196
assert conditional_reaudit["prediction_integrity"]["primary_refits"] == 49
assert conditional_reaudit["prediction_integrity"]["all_strictly_chronological"] is True
assert conditional_reaudit["prediction_integrity"]["all_test_vendors_excluded"] is True
assert conditional_reaudit["prediction_integrity"]["maximum_primary_baseline_difference_b"] < 1e-8
assert conditional_reaudit["prediction_integrity"]["maximum_primary_candidate_difference_b"] < 1e-8
assert "IKP Signal Audit" in workbook.sheetnames
ikp_sheet = workbook["IKP Signal Audit"]
ikp_summary = {
    ikp_sheet.cell(row, 1).value: ikp_sheet.cell(row, 2).value
    for row in range(6, 40)
}
assert ikp_summary["Raw calibration configurations"] == 93
assert ikp_summary["Distinct weight bases after serving collapse"] == 87
assert ikp_summary["Exact overlap models"] == ikp_result["incremental_overlap"]["models"]
assert ikp_summary["Exact overlap families"] == ikp_result["incremental_overlap"]["families"]
assert ikp_summary["Decision-derived evidence-level weight"] == ikp_result["decision"]["incremental_evidence_weight"]
assert ikp_summary["Final Fable weight with 50% crowd"] == ikp_result["decision"]["incremental_final_weight_when_crowd_is_50pct"]
assert math.isclose(
    ikp_summary["Strict Fable estimate (T)"],
    ikp_result["target_signal"]["fable"]["strict_open_only_release_and_vendor_holdout"]["mean"]["estimates"]["forward_inverse"]["estimated_b"] / 1000,
    rel_tol=0,
    abs_tol=1e-12,
)
ikp_values = list(ikp_sheet.iter_rows(values_only=True))
conditional_header = next(
    index for index, row in enumerate(ikp_values)
    if len(row) >= 2 and row[0] == "Benchmark" and row[1] == "Panel bases"
)
conditional_rows = [row for row in ikp_values[conditional_header + 1 : conditional_header + 5]]
assert [row[0] for row in conditional_rows] == ["MMLU", "MMLU-Pro", "GPQA Diamond", "SimpleQA"]
assert conditional_rows[0][2:4] == (16, 11)
assert conditional_rows[0][8:10] == ("3/4", "SUPPORTIVE")
assert conditional_rows[2][2:4] == (18, 9)
assert conditional_rows[2][8:10] == ("4/4", "ROBUST")
narrative_header = next(
    index for index, row in enumerate(ikp_values)
    if len(row) >= 4 and row[0] == "Claim" and row[1] == "Narrative"
)
narrative_rows = [row for row in ikp_values[narrative_header + 1 : narrative_header + 7]]
assert len(narrative_rows) == 6
assert all(row[3] == "STALE" for row in narrative_rows)

assert "Primary Evidence" in workbook.sheetnames
primary_sheet = workbook["Primary Evidence"]
primary_summary = {
    primary_sheet.cell(row, 1).value: primary_sheet.cell(row, 2).value
    for row in range(6, 20)
}
assert primary_summary["Official Sol no-CoT horizon (min)"] == 3.6
assert primary_summary["Official GPT-5.5 comparator (min)"] == 2.3
assert primary_summary["Chronological developer-holdouts"] == 16
assert primary_summary["Held-out developers"] == 5
assert primary_summary["Mapping max / min"] > 6
assert primary_summary["Incremental live weight"] == 0
assert primary_summary["Headline forecasts changed"] == "No"
primary_values = list(primary_sheet.iter_rows(values_only=True))
evidence_header = next(
    index for index, row in enumerate(primary_values)
    if len(row) >= 2 and row[0] == "Evidence ID" and row[1] == "Model"
)
control_header = next(
    index for index, row in enumerate(primary_values)
    if len(row) >= 3 and row[0] == "Record type" and row[1] == "Series" and row[2] == "Model"
)
evidence_rows = [row for row in primary_values[evidence_header + 1 : control_header] if row and row[0] and row[0] != "Chronological holdouts, same-size controls, and Sol mapping sensitivities"]
control_rows = [row for row in primary_values[control_header + 1 :] if row and row[0]]
assert len(evidence_rows) == 5
assert len(control_rows) == 33
assert len({row[0] for row in evidence_rows}) == 5
assert sum(row[5] is not None for row in evidence_rows) == 2
assert sum(row[0] == "chronological_developer_holdout" for row in control_rows) == 16
assert any(row[0] == "exact_open_lineage_control" and row[2] == "Kimi K2.6" for row in control_rows)

primary_result = json.loads(FRONTIER_PRIMARY_AUDIT.read_text(encoding="utf-8"))
assert primary_result["official_measurements"]["gpt_5_6_sol_nocot_minutes"] == 3.6
assert primary_result["official_measurements"]["gpt_5_5_comparator_minutes"] == 2.3
assert primary_result["heldout_backtest"]["incremental_bootstrap"]["ci_90"][0] < 0
assert primary_result["heldout_backtest"]["incremental_bootstrap"]["ci_90"][1] > 0
assert primary_result["sol_mapping_sensitivity"]["nonbaseline_method_max_over_min"] > 6
assert primary_result["decision"]["apply_fable_mythos_shared_weight_identity"] is True
assert primary_result["decision"]["treat_opus_fallback_as_shared_base"] is False
assert primary_result["decision"]["incremental_live_weight"] == 0

assert "No-CoT Date Audit" in workbook.sheetnames
no_cot_date_sheet = workbook["No-CoT Date Audit"]
no_cot_date_summary = {
    no_cot_date_sheet.cell(row, 1).value: no_cot_date_sheet.cell(row, 2).value
    for row in range(6, 18)
}
assert no_cot_date_summary["No-CoT models"] == 49
assert no_cot_date_summary["Models with exact day-level dates"] == 49
assert no_cot_date_summary["Models remaining month-only"] == 0
assert no_cot_date_summary["Explicit date-only overrides"] == 4
assert no_cot_date_summary["Parameter identities added by overrides"] == 0
assert math.isclose(no_cot_date_summary["Paper time-horizon doubling (days)"], 373.0, rel_tol=0, abs_tol=1e-12)
assert math.isclose(no_cot_date_summary["Exact-date adjusted time law (days)"], 365.8606542885574, rel_tol=0, abs_tol=1e-9)
assert math.isclose(no_cot_date_summary["Paper token-horizon doubling (days)"], 437.0, rel_tol=0, abs_tol=1e-12)
assert math.isclose(no_cot_date_summary["Exact-date adjusted token law (days)"], 438.35702998585356, rel_tol=0, abs_tol=1e-9)
assert no_cot_date_summary["No-CoT branch weight changed"] == "No"
no_cot_date_rows = [
    row for row in no_cot_date_sheet.iter_rows(min_row=22, max_row=70, values_only=True)
    if row[0]
]
assert len(no_cot_date_rows) == 49
assert all(isinstance(row[2], datetime) and isinstance(row[3], datetime) for row in no_cot_date_rows)
assert sum(row[6] == "Yes" for row in no_cot_date_rows) == 4
assert all("date_only_no_epoch_parameter_join" in row[7] for row in no_cot_date_rows if row[6] == "Yes")

no_cot_evidence = workbook["No-CoT Evidence"]
frontier_dates = [no_cot_evidence.cell(row, 2).value for row in range(7, 21)]
open_dates = [no_cot_evidence.cell(row, 3).value for row in range(25, 60)]
assert len(frontier_dates) == 14 and len(open_dates) == 35
assert all(isinstance(value, datetime) for value in frontier_dates + open_dates)

horizon_laws = workbook["Horizon Laws"]
assert math.isclose(horizon_laws["B12"].value, 365.8606542885574, rel_tol=0, abs_tol=1e-9)
assert math.isclose(horizon_laws["B15"].value, 438.35702998585356, rel_tol=0, abs_tol=1e-9)
assert math.isclose(horizon_laws["B64"].value, 8.132252384929469, rel_tol=0, abs_tol=1e-10)
assert horizon_laws["B68"].value == "No"
architecture_result = json.loads(NO_COT_ARCHITECTURE_AUDIT.read_text(encoding="utf-8"))
assert architecture_result["inventory"]["models"] == 35
assert architecture_result["paper_relationship_reproduction"]["moe"]["pareto_n"] == 5
assert architecture_result["paired_comparisons"]["chronological_developer_holdout"]["direct_architecture_minus_pooled"]["ci_90"][0] > 0
assert architecture_result["decision"]["replace_pooled_live_elasticity_with_moe_specific"] is False

assert "AA Expansion Audit" in workbook.sheetnames
aa_expansion_sheet = workbook["AA Expansion Audit"]
aa_expansion_summary = {
    aa_expansion_sheet.cell(row, 1).value: aa_expansion_sheet.cell(row, 2).value
    for row in range(6, 29)
}
assert aa_expansion_summary["Current manually curated checkpoints"] == 50
assert aa_expansion_summary["Exact open-weight AA↔Epoch checkpoints"] == 63
assert aa_expansion_summary["Lower-scoring duplicate configurations removed"] == 9
assert aa_expansion_summary["Current/exact overlaps reconciled"] == 19
assert aa_expansion_summary["Expanded unique checkpoints"] == 94
assert aa_expansion_summary["Expanded developers"] == 27
assert aa_expansion_summary["Eligible held-out predictions"] == 46
assert aa_expansion_summary["Incremental expanded-AA weight"] == 0
assert aa_expansion_summary["Change live AA branch"] == "No"
aa_expansion_values = list(aa_expansion_sheet.iter_rows(values_only=True))
scope_header = next(
    index for index, row in enumerate(aa_expansion_values)
    if len(row) >= 2 and row[0] == "Scope" and row[1] == "Tests"
)
overlap_header = next(
    index for index, row in enumerate(aa_expansion_values)
    if len(row) >= 2 and row[0] == "Current model" and row[1] == "Exact AA model"
)
model_headers = [
    index for index, row in enumerate(aa_expansion_values)
    if len(row) >= 3 and row[0] == "Model" and row[1] == "Release" and row[2] == "AA"
]
assert len(model_headers) == 2
frontier_header, aa_panel_header = model_headers
scope_rows = [row for row in aa_expansion_values[scope_header + 1 : frontier_header] if row and row[0] in {"all", "current_panel", "exact_additions", "frontier_like"}]
frontier_rows = [row for row in aa_expansion_values[frontier_header + 1 : overlap_header] if row and len(row) > 6 and row[6] == "Current retained"]
overlap_rows = [row for row in aa_expansion_values[overlap_header + 1 : aa_panel_header] if row and len(row) > 2 and row[2] in {"normalized exact model label", "manual exact checkpoint alias"}]
aa_panel_rows = [row for row in aa_expansion_values[aa_panel_header + 1 :] if row and row[0]]
assert len(scope_rows) == 4
assert len(frontier_rows) == 9
assert all(row[6] == "Current retained" for row in frontier_rows)
assert len(overlap_rows) == 19
assert len({row[0] for row in overlap_rows}) == 19
assert len(aa_panel_rows) == 94
assert len({row[0] for row in aa_panel_rows}) == 94

assert "AA Inference Audit" in workbook.sheetnames
aa_inference_sheet = workbook["AA Inference Audit"]
aa_inference_result = json.loads(AA_INFERENCE_RESULT.read_text(encoding="utf-8"))
aa_inference_data = aa_inference_result["data_audit"]
aa_inference_summary = {
    aa_inference_sheet.cell(row, 1).value: aa_inference_sheet.cell(row, 2).value
    for row in range(6, 30)
}
assert aa_inference_summary["Raw AA model configurations"] == aa_inference_data["raw_models"]
assert aa_inference_summary["Open-weight parameter/score/date configurations"] == aa_inference_data["open_weight_parameter_score_date_configurations"]
assert aa_inference_summary["Unique detailed checkpoint groups"] == aa_inference_data["unique_checkpoint_groups"]
assert aa_inference_summary["Lower-score configurations removed"] == aa_inference_data["lower_score_configurations_removed"]
assert aa_inference_summary["Creators in detailed parameter panel"] == aa_inference_data["creators"]
assert aa_inference_summary["Token-covered checkpoint groups"] == aa_inference_data["token_covered_checkpoint_groups"]
assert aa_inference_summary["Exact AA↔Epoch crosschecks"] == aa_inference_data["epoch_exact_crosschecks"]
assert aa_inference_summary["Crosschecks with visible metadata disagreement"] == aa_inference_data["epoch_crosschecks_with_metadata_disagreement"]
assert aa_inference_summary["All reasoning/non-reasoning pairs"] == aa_inference_result["reasoning_configuration_pairs"]["all"]["pairs"]
assert aa_inference_summary["Strict exact open-weight pairs"] == aa_inference_result["same_weight_reasoning_pairs"]["pairs"]
assert aa_inference_summary["Incremental detailed-panel weight"] == aa_inference_result["decision"]["incremental_detailed_panel_weight"]
assert aa_inference_summary["Incremental inference-budget weight"] == aa_inference_result["decision"]["incremental_inference_budget_weight"]
assert aa_inference_summary["Incremental reasoning-standardization weight"] == aa_inference_result["decision"]["incremental_reasoning_standardization_weight"]
assert aa_inference_summary["Change live AA branch"] == ("Yes" if aa_inference_result["decision"]["change_live_aa_branch"] else "No")
aa_inference_values = list(aa_inference_sheet.iter_rows(values_only=True))
comparison_header = next(
    index for index, row in enumerate(aa_inference_values)
    if len(row) >= 2 and row[0] == "Branch" and row[1] == "Scope"
)
creator_header = next(
    index for index, row in enumerate(aa_inference_values)
    if len(row) >= 2 and row[0] == "Creator" and row[1] == "Pairs"
)
conflict_header = next(
    index for index, row in enumerate(aa_inference_values)
    if len(row) >= 2 and row[0] == "Epoch checkpoint" and row[1] == "Epoch model"
)
pair_header = next(
    index for index, row in enumerate(aa_inference_values)
    if len(row) >= 3 and row[0] == "Creator" and row[1] == "Release" and row[2] == "Open weights"
)
detailed_panel_header = next(
    index for index, row in enumerate(aa_inference_values)
    if len(row) >= 3 and row[0] == "Model" and row[1] == "Creator" and row[2] == "Release"
)
comparison_rows = [row for row in aa_inference_values[comparison_header + 1 : creator_header] if row and row[0] in {f"Detailed {aa_inference_data['unique_checkpoint_groups']} vs current 50", "Measured token budget", "Portable reasoning standardization", "Creator-aware standardization"}]
creator_rows = [row for row in aa_inference_values[creator_header + 1 : conflict_header] if row and row[0] and isinstance(row[1], (int, float))]
conflict_rows = [row for row in aa_inference_values[conflict_header + 1 : pair_header] if row and str(row[0]).startswith("checkpoint:")]
pair_rows = [row for row in aa_inference_values[pair_header + 1 : detailed_panel_header] if row and row[0] and row[0] != f"Complete {aa_inference_data['unique_checkpoint_groups']}-checkpoint detailed parameter panel"]
detailed_panel_rows = [row for row in aa_inference_values[detailed_panel_header + 1 :] if row and row[0]]
assert len(comparison_rows) == 8
assert len(creator_rows) == len(aa_inference_result["reasoning_configuration_pairs"]["all"]["creator_medians"])
assert len(conflict_rows) == aa_inference_data["epoch_crosschecks_with_metadata_disagreement"]
assert len(pair_rows) == aa_inference_result["reasoning_configuration_pairs"]["all"]["pairs"]
assert len(detailed_panel_rows) == aa_inference_data["unique_checkpoint_groups"]
assert len({(row[0], row[1], row[5], row[8]) for row in pair_rows}) == len(pair_rows)
assert len({(row[0], row[1], row[2]) for row in detailed_panel_rows}) == len(detailed_panel_rows)

assert "AA Operational Audit" in workbook.sheetnames
aa_operational_sheet = workbook["AA Operational Audit"]
aa_operational_result = json.loads(AA_OPERATIONAL_RESULT.read_text(encoding="utf-8"))
aa_operational_data = aa_operational_result["data_audit"]
aa_operational_summary = {
    aa_operational_sheet.cell(row, 1).value: aa_operational_sheet.cell(row, 2).value
    for row in range(6, 27)
}
assert aa_operational_summary["Raw AA model configurations"] == aa_operational_data["raw_model_configurations"]
assert aa_operational_summary["Deduplicated open parameter checkpoints"] == aa_operational_data["deduplicated_open_parameter_checkpoints"]
assert aa_operational_summary["Price-covered checkpoints"] == aa_operational_data["coverage"]["blended_price"]["checkpoints"]
assert aa_operational_summary["Price-covered developers"] == aa_operational_data["coverage"]["blended_price"]["developers"]
assert aa_operational_summary["Speed-covered checkpoints"] == aa_operational_data["coverage"]["output_speed"]["checkpoints"]
assert aa_operational_summary["Speed-covered developers"] == aa_operational_data["coverage"]["output_speed"]["developers"]
assert aa_operational_summary["Exact AA↔OpenRouter checkpoints"] == 28
assert aa_operational_summary["Cross-source price Spearman"] > 0.8
assert aa_operational_summary["Cross-source raw-speed Spearman"] < aa_operational_summary["Cross-source price Spearman"] - 0.3
assert aa_operational_summary["Incremental AA operational price weight"] == 0
assert aa_operational_summary["Incremental speed weight"] == 0
assert aa_operational_summary["Incremental latency weight"] == 0
assert aa_operational_summary["Change live price weight"] == "No"
aa_operational_values = list(aa_operational_sheet.iter_rows(values_only=True))
operational_comparison_header = next(
    index for index, row in enumerate(aa_operational_values)
    if len(row) >= 2 and row[0] == "Branch" and row[1] == "Scope"
)
operational_crosscheck_header = next(
    index for index, row in enumerate(aa_operational_values)
    if len(row) >= 2 and row[0] == "Canonical checkpoint" and row[1] == "Epoch model"
)
operational_panel_header = next(
    index for index, row in enumerate(aa_operational_values)
    if len(row) >= 3 and row[0] == "Model" and row[1] == "Creator" and row[2] == "Release"
)
operational_comparison_rows = [
    row for row in aa_operational_values[operational_comparison_header + 1 : operational_crosscheck_header]
    if row and row[0] in {
        "Price — global", "Price + source regime", "Price — provider median",
        "Price — first party", "Output speed", "Latency / TTFC",
        "Speed + latency", "Cost per task — exploratory", "Time per task — exploratory",
    }
]
operational_crosscheck_rows = [
    row for row in aa_operational_values[operational_crosscheck_header + 1 : operational_panel_header]
    if row and str(row[0]).startswith("checkpoint:")
]
operational_panel_rows = [row for row in aa_operational_values[operational_panel_header + 1 :] if row and row[0]]
assert len(operational_comparison_rows) == 18
assert len(operational_crosscheck_rows) == 28
assert len({row[0] for row in operational_crosscheck_rows}) == 28
assert len(operational_panel_rows) == aa_operational_data["deduplicated_open_parameter_checkpoints"]
assert len({(row[0], row[1], row[2]) for row in operational_panel_rows}) == aa_operational_data["deduplicated_open_parameter_checkpoints"]

assert "Active Param Audit" in workbook.sheetnames
active_transport_sheet = workbook["Active Param Audit"]
active_architecture = {
    active_transport_sheet.cell(row, 1).value: active_transport_sheet.cell(row, 2).value
    for row in range(6, 19)
}
assert active_architecture["K3 total parameters (B)"] == 2780
assert active_architecture["K3 activated parameters (B)"] == 104.2
assert math.isclose(active_architecture["K3 activated fraction (%)"], 100 * 104.2 / 2780, rel_tol=0, abs_tol=1e-12)
assert math.isclose(active_architecture["K3 total / activated ratio"], 2780 / 104.2, rel_tol=0, abs_tol=1e-12)
assert active_architecture["K3 selected routed experts/token"] == 16
assert active_architecture["K3 routed experts"] == 896
assert active_architecture["K3 shared experts"] == 2
assert math.isclose(active_architecture["K3 selected routed-expert fraction (%)"], 100 * 16 / 896, rel_tol=0, abs_tol=1e-12)
assert active_architecture["K3 transformer layers"] == 93
assert active_architecture["K2 total parameters (B)"] == 1040
assert active_architecture["K2 activated parameters (B)"] == 32.6
assert math.isclose(active_architecture["K2 total / activated ratio"], 1040 / 32.6, rel_tol=0, abs_tol=1e-12)

active_result = json.loads(ACTIVE_TRANSPORT_RESULT.read_text(encoding="utf-8"))
active_values = list(active_transport_sheet.iter_rows(values_only=True))
active_comparison_header = next(
    index for index, row in enumerate(active_values)
    if len(row) >= 2 and row[0] == "Comparison" and row[1] == "Tests"
)
active_target_header = next(
    index for index, row in enumerate(active_values)
    if len(row) >= 3 and row[0] == "Model" and row[1] == "Release" and row[2] == "AA"
)
active_prediction_header = next(
    index for index, row in enumerate(active_values)
    if len(row) >= 3 and row[0] == "Release" and row[1] == "Model" and row[2] == "Developer"
)
active_fold_label = f"Predict active vs total on identical {active_result['inventory']['chronological_predictions']}-checkpoint panel"
active_comparison_rows = [row for row in active_values[active_comparison_header + 1 : active_target_header] if row and row[0] in {active_fold_label, "Convert active→total for ≥15× sparsity"}]
active_target_rows = [row for row in active_values[active_target_header + 1 : active_prediction_header] if row and row[0] in {"Claude Fable 5", "GPT-5.6 Sol", "Kimi K3"}]
active_prediction_rows = [row for row in active_values[active_prediction_header + 1 :] if row and row[0]]
assert len(active_comparison_rows) == 2
assert len(active_target_rows) == 3
assert len(active_prediction_rows) == active_result["inventory"]["chronological_predictions"]
assert len({(row[0], row[1], row[2]) for row in active_prediction_rows}) == len(active_prediction_rows)
active_targets_by_model = {row[0]: row for row in active_target_rows}
active_result_targets = {row["model"]: row for row in active_result["target_sensitivity"]}
assert math.isclose(active_targets_by_model["Claude Fable 5"][5], active_result_targets["Claude Fable 5"]["k3_anchored_total_t"], rel_tol=0, abs_tol=1e-9)
assert math.isclose(active_targets_by_model["GPT-5.6 Sol"][5], active_result_targets["GPT-5.6 Sol"]["k3_anchored_total_t"], rel_tol=0, abs_tol=1e-9)
assert math.isclose(active_targets_by_model["Kimi K3"][4], 104.2, rel_tol=0, abs_tol=1e-9)
assert math.isclose(active_targets_by_model["Kimi K3"][5], 2.78, rel_tol=0, abs_tol=1e-9)

assert active_result["inventory"]["active_parameter_checkpoints"] > 0
assert active_result["inventory"]["chronological_predictions"] == len(active_prediction_rows)
assert active_result["compute_branch_independence"]["target_models_with_epoch_training_compute_estimate"] == 1
assert active_result["compute_branch_independence"]["target_models_with_disclosed_training_compute"] == 0
assert active_result["compute_branch_independence"]["independent_target_evidence"] is False
assert active_result["decision"]["incremental_live_weight"] == 0
assert active_result["decision"]["change_headline_forecasts"] is False

formula_workbook = load_workbook(WORKBOOK, data_only=False, read_only=True)
active_formula_sheet = formula_workbook["Active Param Audit"]
assert math.isclose(active_formula_sheet["B8"].value, 100 * 104.2 / 2780, rel_tol=0, abs_tol=1e-12)
assert math.isclose(active_formula_sheet["B9"].value, 2780 / 104.2, rel_tol=0, abs_tol=1e-12)
assert active_formula_sheet["F30"].value == "=$B$6*E30/$B$7/1000"
formula_workbook.close()

assert "ECI Component Audit" in workbook.sheetnames
eci_component_sheet = workbook["ECI Component Audit"]
eci_component_summary = {
    eci_component_sheet.cell(row, 1).value: eci_component_sheet.cell(row, 2).value
    for row in range(6, 22)
}
assert eci_component_summary["Expanded total-parameter checkpoints"] == 84
assert eci_component_summary["Exact open-weight Epoch additions"] == 26
assert eci_component_summary["Active-parameter checkpoints"] == 89
assert eci_component_summary["Eligible benchmark comparisons"] == 13
assert eci_component_summary["Components supported after correction"] == 0
assert eci_component_summary["Incremental component weight"] == 0
assert eci_component_summary["Change live ECI center"] == "No"
eci_component_values = list(eci_component_sheet.iter_rows(values_only=True))
benchmark_header = next(
    index for index, row in enumerate(eci_component_values)
    if len(row) >= 2 and row[0] == "Benchmark" and row[1] == "Coverage"
)
panel_header = next(
    index for index, row in enumerate(eci_component_values)
    if len(row) >= 2 and row[0] == "Model" and row[1] == "Release"
)
benchmark_rows = [row for row in eci_component_values[benchmark_header + 1 : panel_header] if row and row[0] and row[0] != "Complete 84-checkpoint total-parameter admission ledger"]
panel_rows = [row for row in eci_component_values[panel_header + 1 :] if row and row[0]]
assert len(benchmark_rows) == 13
assert all(row[12] == "Not supported" for row in benchmark_rows)
assert len(panel_rows) == 84
assert len({row[0] for row in panel_rows}) == 84

assert "ECI Multivariate Audit" in workbook.sheetnames
eci_multivariate_sheet = workbook["ECI Multivariate Audit"]
eci_multivariate_summary = {
    eci_multivariate_sheet.cell(row, 1).value: eci_multivariate_sheet.cell(row, 2).value
    for row in range(6, 19)
}
assert eci_multivariate_summary["Parameter-map checkpoints"] == 89
assert eci_multivariate_summary["Developer families"] == 40
assert eci_multivariate_summary["Narrow ECI CI checkpoints"] == 41
assert eci_multivariate_summary["Broad ECI CI checkpoints"] == 48
assert eci_multivariate_summary["Unique component measurements"] == 723
assert eci_multivariate_summary["Component benchmarks"] == 50
assert eci_multivariate_summary["Primary outer predictions"] == 73
assert eci_multivariate_summary["Narrow-ECI-CI-only predictions"] == 28
assert eci_multivariate_summary["Incremental live weight"] == 0
assert eci_multivariate_summary["Change headline forecasts"] == "No"

eci_multivariate_values = list(eci_multivariate_sheet.iter_rows(values_only=True))
target_header = next(
    index for index, row in enumerate(eci_multivariate_values)
    if len(row) >= 3 and row[0] == "Model" and row[1] == "Release" and row[2] == "Full train"
)
coverage_header = next(
    index for index, row in enumerate(eci_multivariate_values)
    if len(row) >= 3 and row[0] == "Benchmark" and row[1] == "Models" and row[2] == "Families"
)
prediction_headers = [
    index for index, row in enumerate(eci_multivariate_values)
    if len(row) >= 3 and row[0] == "Release" and row[1] == "Model" and row[2] == "Family"
]
assert len(prediction_headers) == 2
primary_header, narrow_ci_header = prediction_headers
target_rows = [
    row for row in eci_multivariate_values[target_header + 1 : coverage_header]
    if row and row[0] in {"Claude Fable 5", "GPT-5.6 Sol"}
]
coverage_rows = [
    row for row in eci_multivariate_values[coverage_header + 1 : primary_header]
    if row and row[0] and isinstance(row[1], (int, float))
]
primary_rows = [
    row for row in eci_multivariate_values[primary_header + 1 : narrow_ci_header]
    if row and isinstance(row[0], datetime)
]
narrow_ci_rows = [
    row for row in eci_multivariate_values[narrow_ci_header + 1 :]
    if row and isinstance(row[0], datetime)
]
assert len(target_rows) == 2
assert len(coverage_rows) == 50
assert sum(row[1] for row in coverage_rows) == 723
assert len(primary_rows) == 73
assert len(narrow_ci_rows) == 28
assert all(row[7] < row[0] for row in primary_rows + narrow_ci_rows)
assert all(row[3] is False for row in narrow_ci_rows)
targets_by_model = {row[0]: row for row in target_rows}
assert targets_by_model["Claude Fable 5"][5] == 1
assert targets_by_model["Claude Fable 5"][14] is False
assert targets_by_model["GPT-5.6 Sol"][5] == 2

multivariate_result = json.loads(ECI_MULTIVARIATE_RESULT.read_text(encoding="utf-8"))
assert multivariate_result["backtest"]["total"]["all"]["candidate"]["median_multiplicative_error"] < multivariate_result["backtest"]["total"]["all"]["baseline"]["median_multiplicative_error"]
assert multivariate_result["backtest"]["total"]["all"]["paired_family_bootstrap"]["ci_90"][1] > 0
assert multivariate_result["narrow_eci_ci_only_training_backtest"]["total"]["paired_family_bootstrap"]["ci_90"][1] > 0
assert multivariate_result["decision"]["promote_multivariate_component_branch"] is False

multivariate_formula_workbook = load_workbook(WORKBOOK, data_only=False, read_only=True)
multivariate_formula_sheet = multivariate_formula_workbook["ECI Multivariate Audit"]
first_target_excel_row = target_header + 2
assert multivariate_formula_sheet.cell(first_target_excel_row, 9).value == f"=H{first_target_excel_row}/G{first_target_excel_row}"
assert multivariate_formula_sheet.cell(first_target_excel_row, 14).value == f"=M{first_target_excel_row}/L{first_target_excel_row}"
multivariate_formula_workbook.close()

assert "Post-training Audit" in workbook.sheetnames
posttraining_sheet = workbook["Post-training Audit"]
posttraining_result = json.loads(POSTTRAINING_LINEAGE_RESULT.read_text(encoding="utf-8"))
posttraining_inventory = posttraining_result["inventory"]
posttraining_summary = {
    posttraining_sheet.cell(row, 1).value: posttraining_sheet.cell(row, 2).value
    for row in range(6, 20)
}
assert posttraining_summary["Epoch base-model links"] == posttraining_inventory["base_model_links"]
assert posttraining_summary["Unique exact parent matches"] == posttraining_inventory["unique_exact_parent_matches"]
assert posttraining_summary["Same-parameter links (±1%)"] == posttraining_inventory["same_parameter_links_1pct"]
assert posttraining_summary["Same-parameter links, both open"] == posttraining_inventory["same_parameter_both_open_links"]
assert posttraining_summary["Open language candidates"] == posttraining_inventory["candidate_open_language_same_parameter_links"]
assert posttraining_summary["Measured lineage edges"] == posttraining_inventory["admitted_measured_lineage_edges"]
assert posttraining_summary["Measured lineage bases"] == posttraining_inventory["admitted_measured_lineage_bases"]
assert posttraining_summary["Measured developers"] == posttraining_inventory["admitted_developers"]
assert posttraining_summary["Matched measurements"] == posttraining_inventory["matched_measurements"]
assert posttraining_summary["Matched ECI components"] == posttraining_inventory["matched_component_measurements"]
assert posttraining_summary["Edges with finetune compute"] == posttraining_inventory["edges_with_finetune_compute"]
assert posttraining_summary["No-CoT / METR lineage edges"] == f"{posttraining_inventory['nocot_lineage_edges']} / {posttraining_inventory['metr_lineage_edges']}"
assert posttraining_summary["Incremental live weight"] == posttraining_result["decision"]["incremental_live_weight"]
assert posttraining_summary["Change headline forecasts"] == ("Yes" if posttraining_result["decision"]["change_headline_forecasts"] else "No")

posttraining_values = list(posttraining_sheet.iter_rows(values_only=True))
posttraining_evidence_header = next(
    index for index, row in enumerate(posttraining_values)
    if len(row) >= 3 and row[0] == "Lineage" and row[1] == "Claim" and row[2] == "Evidence grade"
)
posttraining_edge_header = next(
    index for index, row in enumerate(posttraining_values)
    if len(row) >= 3 and row[0] == "Child" and row[1] == "Parent" and row[2] == "Child date"
)
posttraining_prediction_header = next(
    index for index, row in enumerate(posttraining_values)
    if len(row) >= 3 and row[0] == "Signal" and row[1] == "Child" and row[2] == "Parent"
)
posttraining_frontier_header = next(
    index for index, row in enumerate(posttraining_values)
    if len(row) >= 4 and row[0] == "Chain" and row[1] == "Mode" and row[3] == "Model"
)
posttraining_evidence_rows = [
    row for row in posttraining_values[posttraining_evidence_header + 1 : posttraining_edge_header]
    if row and row[0] and row[0] != "Exact open-weight same-parameter lineage edges"
]
posttraining_edge_rows = [
    row for row in posttraining_values[posttraining_edge_header + 1 : posttraining_prediction_header]
    if row and row[0] and row[0] != "Strictly earlier, endpoint-group-held-out lineage predictions"
]
posttraining_prediction_rows = [
    row for row in posttraining_values[posttraining_prediction_header + 1 : posttraining_frontier_header]
    if row and row[0] and row[0] != "Proprietary asserted-same-base sensitivity — diagnostic only, not actual parameter ratios"
]
posttraining_frontier_rows = [
    row for row in posttraining_values[posttraining_frontier_header + 1 :]
    if row and row[0]
]
assert len(posttraining_evidence_rows) == 4
assert len(posttraining_edge_rows) == posttraining_inventory["admitted_measured_lineage_edges"]
assert len(posttraining_prediction_rows) == posttraining_inventory["eci_prediction_edges"] + posttraining_inventory["aa_prediction_edges"]
assert len(posttraining_frontier_rows) == 18
assert len({(row[0], row[1], row[2], row[3]) for row in posttraining_edge_rows}) == len(posttraining_edge_rows)
assert all(row[3] <= row[2] for row in posttraining_edge_rows)
assert all(isinstance(row[12], datetime) for row in posttraining_prediction_rows)
assert all(row[13] is True for row in posttraining_prediction_rows)
assert all(row[15] == "user_asserted_not_publicly_disclosed" for row in posttraining_frontier_rows)
assert any(row[14] is True for row in posttraining_frontier_rows)

assert posttraining_result["lineage_backtests"]["eci"]["collapsed_vs_baseline"]["ci_90"][1] > 0
assert posttraining_result["lineage_backtests"]["aa"]["collapsed_vs_baseline"]["ci_90"][1] < 0
assert posttraining_result["promotion_gates"]["aa_signal_bases_at_least_6"] is False
assert posttraining_result["decision"]["promote_posttraining_correction"] is False
assert posttraining_result["decision"]["incremental_live_weight"] == 0

assert "Active Price Audit" in workbook.sheetnames
active_price_sheet = workbook["Active Price Audit"]
active_price_summary = {
    active_price_sheet.cell(row, 1).value: active_price_sheet.cell(row, 2).value
    for row in range(6, 26)
}
assert active_price_summary["Active-parameter labels"] == 63
assert active_price_summary["AA disclosed active labels"] == 45
assert active_price_summary["Dense config controls"] == 18
assert active_price_summary["Matched developers"] == 15
assert active_price_summary["Release-ordered predictions"] == 45
assert active_price_summary["Prediction developers"] == 11
assert active_price_summary["High-sparsity transport tests"] == 16
assert active_price_summary["High-sparsity developers"] == 7
assert active_price_summary["Transport candidate median error (×)"] < active_price_summary["Direct-total baseline median error (×)"]
assert active_price_summary["Performance gate passed"] == "No"
assert active_price_summary["Coverage gate passed"] == "No"
assert active_price_summary["Incremental live weight"] == 0
active_price_values = list(active_price_sheet.iter_rows(values_only=True))
active_price_target_header = next(index for index, row in enumerate(active_price_values) if len(row) >= 5 and row[3] == "Model" and row[4] == "Release")
active_price_match_header = next(index for index, row in enumerate(active_price_values) if len(row) >= 2 and row[0] == "Epoch checkpoint" and row[1] == "Epoch model")
active_price_prediction_header = next(index for index, row in enumerate(active_price_values) if len(row) >= 3 and row[0] == "Checkpoint" and row[1] == "Release" and row[2] == "Model")
active_price_target_rows = [row for row in active_price_values[active_price_target_header + 1 : active_price_match_header] if len(row) > 3 and row[3] in {"Claude Fable 5", "GPT-5.6 Sol", "Kimi K3"}]
active_price_match_rows = [row for row in active_price_values[active_price_match_header + 1 : active_price_prediction_header] if row and row[0] and row[0] != "Complete release-ordered developer-held-out prediction ledger"]
active_price_prediction_rows = [row for row in active_price_values[active_price_prediction_header + 1 :] if row and row[0]]
assert len(active_price_target_rows) == 3
assert len(active_price_match_rows) == 93
assert len({row[0] for row in active_price_match_rows}) == 93
dense_config_rows = [row for row in active_price_match_rows if row[11] == "dense_config_active_equals_total"]
assert len(dense_config_rows) == 18
assert all(row[7] == row[4] for row in dense_config_rows)
assert all(row[10] == "dense_config" for row in dense_config_rows)
assert len(active_price_prediction_rows) == 45
assert len({row[0] for row in active_price_prediction_rows}) == 45
active_price_targets_by_model = {row[3]: row for row in active_price_target_rows}
with OPENROUTER_ACTIVE_PRICE_TARGETS.open(newline="", encoding="utf-8-sig") as handle:
    expected_active_price_targets = {
        row["model"]: row for row in csv.DictReader(handle)
    }
for model in ("Claude Fable 5", "GPT-5.6 Sol", "Kimi K3"):
    assert math.isclose(
        active_price_targets_by_model[model][9],
        float(
            expected_active_price_targets[model][
                "k3_anchored_total_score_date_price_t"
            ]
        ),
        rel_tol=0,
        abs_tol=1e-9,
    )

assert "OpenRouter Audit" in workbook.sheetnames
openrouter_sheet = workbook["OpenRouter Audit"]
summary_values = {
    openrouter_sheet.cell(row, 1).value: openrouter_sheet.cell(row, 2).value
    for row in range(6, 15)
}
collection_audit = json.loads(OPENROUTER_COLLECTION_AUDIT.read_text(encoding="utf-8"))
assert summary_values["OpenRouter text models retained"] == collection_audit["eligible_text_model_count"]
assert summary_values["Manually matched OpenRouter aliases"] == 95
assert summary_values["Unique Epoch calibration checkpoints"] == 93
assert summary_values["Tok/s adds robust incremental information"] == "No"
openrouter_values = list(openrouter_sheet.iter_rows(values_only=True))
match_header_row = next(
    index for index, row in enumerate(openrouter_values)
    if len(row) >= 2 and row[0] == "OpenRouter ID" and row[1] == "OpenRouter name"
)
match_rows = [row for row in openrouter_values[match_header_row + 1 :] if row and row[0]]
assert len(match_rows) == collection_audit["eligible_text_model_count"]
assert len({row[0] for row in match_rows}) == collection_audit["eligible_text_model_count"]

assert "OR Time Stability" in workbook.sheetnames
openrouter_time_sheet = workbook["OR Time Stability"]
openrouter_time_summary = {
    openrouter_time_sheet.cell(row, 1).value: openrouter_time_sheet.cell(row, 2).value
    for row in range(6, 22)
}
temporal_result = json.loads(OPENROUTER_TEMPORAL_RESULT.read_text(encoding="utf-8"))
assert openrouter_time_summary["Immutable refresh snapshots"] == temporal_result["inventory"]["immutable_snapshots"]
assert openrouter_time_summary["Current lossless daily rows"] == temporal_result["inventory"]["current_daily_rows"]
assert openrouter_time_summary["Daily rows across snapshot history"] == temporal_result["inventory"]["history_daily_rows"]
assert openrouter_time_summary["Calibration checkpoints changed by tier split"] == 0
assert openrouter_time_summary["Recommended incremental tok/s weight"] == 0
assert openrouter_time_summary["Change live forecast"] == "No"
openrouter_time_values = list(openrouter_time_sheet.iter_rows(values_only=True))
history_header = next(index for index, row in enumerate(openrouter_time_values) if row and row[0] == "Snapshot ID")
model_header = next(index for index, row in enumerate(openrouter_time_values) if row and row[0] == "OpenRouter ID" and len(row) > 2 and row[2] == "Days")
history_rows = [row for row in openrouter_time_values[history_header + 1 : model_header] if row and row[0] and row[0] != "Model-level daily stability sample — most volatile plus all frontier targets"]
assert len(history_rows) == temporal_result["inventory"]["immutable_snapshots"]
assert len({row[0] for row in history_rows}) == len(history_rows)

assert "OR Price Audit" in workbook.sheetnames
openrouter_price_sheet = workbook["OR Price Audit"]
openrouter_price_summary = {
    openrouter_price_sheet.cell(row, 1).value: openrouter_price_sheet.cell(row, 2).value
    for row in range(6, 20)
}
official_audit = json.loads(OPENROUTER_OFFICIAL_AUDIT.read_text(encoding="utf-8"))
collection_audit = json.loads(OPENROUTER_COLLECTION_AUDIT.read_text(encoding="utf-8"))
assert openrouter_price_summary["Endpoint/service-tier rows"] == collection_audit["endpoint_tier_row_count"]
assert openrouter_price_summary["Rows with high-context prices"] == collection_audit["endpoint_tier_rows_with_high_context_price"]
assert math.isclose(openrouter_price_summary["Official price exact share"], official_audit["official_price_row_exact_share"], rel_tol=0, abs_tol=1e-12)
assert openrouter_price_summary["Incremental forecast weight"] == 0
openrouter_price_values = list(openrouter_price_sheet.iter_rows(values_only=True))
focal_tier_header = next(index for index, row in enumerate(openrouter_price_values) if len(row) >= 5 and row[0] == "OpenRouter ID" and row[4] == "Tier")
mismatch_header = next(index for index, row in enumerate(openrouter_price_values) if len(row) >= 3 and row[0] == "OpenRouter ID" and row[1] == "Provider tag" and row[2] == "Status")
focal_tier_rows = [row for row in openrouter_price_values[focal_tier_header + 1 : mismatch_header] if row and row[0] and row[0] != "All non-exact official API ↔ frontend price groups"]
mismatch_rows = [row for row in openrouter_price_values[mismatch_header + 1 :] if row and row[0]]
assert len(focal_tier_rows) > 30
assert len(mismatch_rows) == sum(value for key, value in official_audit["comparison_group_counts"].items() if key != "exact")

print(
    {
        "active_forecasts": len(active),
        "contributors": len({row["contributor"] for row in active}),
        "fable_n": counts["Claude Fable 5"],
        "fable_center_t": centers["Claude Fable 5"],
        "sol_n": counts["GPT-5.6 Sol"],
        "sol_center_t": centers["GPT-5.6 Sol"],
        "status": "PASS",
    }
)
