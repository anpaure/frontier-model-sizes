from __future__ import annotations

import csv
import json
import re
import unittest
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "019f6c42-2d53-7743-ab07-6293e2618dd7"
LEDGER = ROOT / "sources/human_parameter_forecasts_2026-07-17.csv"
SITE_DATA = ROOT / "site/public/data/forecast-model.json"
FREEZE = ROOT / "forecast_freezes/2026-07-31-frontier-parameters-v1/forecast_freeze.json"
PRIVACY_RECORD = ROOT / "forecast_freezes/2026-07-31-frontier-parameters-v1/privacy_redaction_2026-08-01.json"
REGISTRY = OUT / "frontier_parameter_prediction_registry_v2.1_2026-07-17.docx"
WORKBOOK = OUT / "frontier_parameter_model_crowd_50pct_2026-07-17.xlsx"
RESPONDENT = re.compile(r"^Respondent R\d{2}$")
FORECAST_ID = re.compile(r"^r\d{2}-[a-z0-9-]+-\d{8}$")


class PollAnonymizationTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        with LEDGER.open(newline="", encoding="utf-8-sig") as handle:
            cls.rows = list(csv.DictReader(handle))
        cls.site = json.loads(SITE_DATA.read_text(encoding="utf-8"))
        cls.freeze = json.loads(FREEZE.read_text(encoding="utf-8"))

    def test_authoritative_ledger_has_only_stable_anonymous_ids(self) -> None:
        self.assertEqual(len(self.rows), 42)
        respondents = {row["contributor"] for row in self.rows}
        self.assertEqual(respondents, {f"Respondent R{i:02d}" for i in range(1, 22)})
        self.assertTrue(all(RESPONDENT.fullmatch(row["contributor"]) for row in self.rows))
        self.assertTrue(all(FORECAST_ID.fullmatch(row["forecast_id"]) for row in self.rows))
        self.assertEqual(len({row["forecast_id"] for row in self.rows}), len(self.rows))
        self.assertTrue(
            all(
                row["provenance"]
                in {
                    "Direct statement from Respondent R01 in this Codex task.",
                    "Relayed by project owner; source not attached.",
                }
                for row in self.rows
            )
        )

    def test_generated_site_and_freeze_expose_no_named_poll_labels(self) -> None:
        for model in self.site["models"]:
            self.assertTrue(
                all(RESPONDENT.fullmatch(value) for value in model["crowd"]["contributors"])
            )
        records = [
            record
            for target in self.freeze["targets"]
            for record in target["crowd_pool"]["records"]
        ]
        self.assertTrue(all(RESPONDENT.fullmatch(row["contributor"]) for row in records))
        self.assertTrue(all(FORECAST_ID.fullmatch(row["forecast_id"]) for row in records))
        self.assertFalse(self.freeze["privacy_redaction"]["name_to_id_mapping_retained"])
        privacy = json.loads(PRIVACY_RECORD.read_text(encoding="utf-8"))
        self.assertFalse(privacy["name_to_id_mapping_retained"])
        self.assertFalse(self.freeze["privacy_redaction"]["numerical_values_changed"])

    def test_current_registry_is_anonymous_and_neutral_metadata(self) -> None:
        document = Document(REGISTRY)
        respondent_cells = [row.cells[0].text.strip() for row in document.tables[0].rows[1:]]
        self.assertEqual(len(respondent_cells), 21)
        self.assertTrue(all(RESPONDENT.fullmatch(value) for value in respondent_cells))
        self.assertEqual(document.core_properties.author, "Project team")
        self.assertEqual(document.core_properties.last_modified_by, "Project team")
        self.assertIn("Anonymous respondents: 21", "\n".join(p.text for p in document.paragraphs))

    def test_publishable_workbook_has_anonymous_note_and_relative_local_paths(self) -> None:
        workbook = load_workbook(WORKBOOK, data_only=False, read_only=True)
        note = workbook["Horizon Estimates"]["O9"].value
        self.assertIn("Respondent R17", note)
        manifest = workbook["Source Manifest"]
        for row in manifest.iter_rows(min_row=6, values_only=True):
            path = row[1]
            if isinstance(path, str) and not path.startswith(("http://", "https://")):
                self.assertFalse(path.startswith("/"), path)
                self.assertNotIn("\\Users\\", path)

    def test_obsolete_named_registry_exports_are_absent(self) -> None:
        old_docx = [
            path
            for path in OUT.glob("frontier_parameter_prediction_registry_v*_2026-07-17.docx")
            if path != REGISTRY
        ]
        self.assertEqual(old_docx, [])
        self.assertEqual(list((ROOT / "qa").glob("prediction-registry*")), [])

    def test_current_ooxml_packages_have_no_unexpected_respondent_labels(self) -> None:
        for path in (REGISTRY, WORKBOOK):
            with zipfile.ZipFile(path) as archive:
                text = "\n".join(
                    archive.read(name).decode("utf-8", errors="ignore")
                    for name in archive.namelist()
                    if name.endswith((".xml", ".rels"))
                )
            for match in re.findall(r"Respondent[^<\n]{0,40}", text):
                if (
                    "Respondents" in match
                    or "respondents" in match
                    or match.startswith("Respondent ID")
                ):
                    continue
                self.assertRegex(match, r"Respondent R\d{2}")


if __name__ == "__main__":
    unittest.main(verbosity=2)
